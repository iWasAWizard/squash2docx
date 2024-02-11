#!/usr/bin/python3

import requests
from docx import Document
from docx.shared import Inches

api_endpoint = "https://api.squash.com/v1/test_cases"

response = requests.get(api_endpoint)
data = response.json()
test_steps = data["test_steps"]
test_cases = data["test_cases"]
procedures = data["procedures"]

document = Document()
steps_table = document.add_table(rows=1, cols=3)
cases_table = document.add_table(rows=1, cols=3)
procs_table = document.add_table(rows=1, cols=3)

for test_step in test_steps:
    steps_table.add_row()
    steps_table.add_cell(test_step["name"])
    steps_table.add_cell(test_step["description"])
    steps_table.add_cell(test_step["expected_result"])

for test_case in test_cases:
    cases_table.add_row()
    cases_table.add_cell(test_case["name"])
    cases_table.add_cell(test_case["description"])
    cases_table.add_cell(test_case["expected_result"])

for procedure in procedures:
    procs_table.add_row()
    procs_table.add_cell(procedure["name"])
    procs_table.add_cell(procedure["description"])
    procs_table.add_cell(procedure["expected_result"])

document.save("squash_tables.docx")

